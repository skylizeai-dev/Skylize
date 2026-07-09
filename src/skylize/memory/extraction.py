"""File → text extraction for org knowledge ingestion (Phase 7b-i formats).

Supported: .md .txt .csv (native), .pdf (pypdf), .docx (python-docx),
.xlsx (openpyxl). Anything else raises UnsupportedFormatError with an honest
message — image OCR (.png/.jpg), .pptx and video are explicitly not yet
supported and must never be silently swallowed.
"""

from __future__ import annotations

import io
from pathlib import PurePosixPath

SUPPORTED_EXTENSIONS = {".md", ".txt", ".csv", ".pdf", ".docx", ".xlsx"}


class UnsupportedFormatError(ValueError):
    def __init__(self, extension: str) -> None:
        super().__init__(
            f"File format '{extension}' is not supported yet. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )
        self.extension = extension


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded file, by extension."""
    ext = PurePosixPath(filename.lower()).suffix
    if ext in {".md", ".txt", ".csv"}:
        return data.decode("utf-8", errors="replace")
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".xlsx":
        return _extract_xlsx(data)
    raise UnsupportedFormatError(ext or "(none)")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ── department routing heuristic ─────────────────────────────────────────
# Mirrors the console's dept taxonomy (website gen_agent_network_data.js).
# Simple keyword scoring — full semantic routing is a later concern.

_DEPT_KEYWORDS: dict[str, tuple[str, ...]] = {
    "marketing": ("marketing", "campaign", "seo", "newsletter", "brand", "ads", "growth"),
    "creative": ("copy", "creative", "design", "video", "content", "caption"),
    "finance": ("invoice", "budget", "revenue", "forecast", "expense", "payment", "p&l", "cash"),
    "sales": ("lead", "deal", "pipeline", "crm", "prospect", "quota", "account"),
    "customer_success": ("ticket", "support", "customer", "churn", "csat", "retention"),
    "operations": ("logistics", "fulfillment", "inventory", "supply", "vendor sla", "process"),
    "engineering": ("api", "deploy", "infrastructure", "backend", "frontend", "devops", "code"),
    "data": ("analytics", "metric", "dashboard", "warehouse", "model", "dataset"),
    "security": ("security", "compliance", "audit", "vulnerability", "incident", "access control"),
    "legal": ("contract", "privacy", "gdpr", "terms", "liability", "nda"),
    "people": ("hiring", "onboarding", "training", "performance review", "playbook"),
    "procurement": ("procurement", "sourcing", "supplier", "purchase order"),
    "product": ("roadmap", "feature", "user research", "experiment", "backlog"),
    "strategy": ("competitor", "market analysis", "expansion", "m&a", "strategy"),
}


def infer_department(text: str) -> str:
    """Best-effort department tag for a document; 'executive_office' if unclear."""
    sample = text[:20_000].lower()
    best, best_score = "executive_office", 0
    for dept, keywords in _DEPT_KEYWORDS.items():
        score = sum(sample.count(k) for k in keywords)
        if score > best_score:
            best, best_score = dept, score
    return best
