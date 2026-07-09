"""DOCX exporter — line-by-line markdown → python-docx Document."""

from __future__ import annotations

import io
import re

from docx import Document

from .base import sanitize_filename

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")


class DocxExporter:
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    def export(self, content_markdown: str, title: str) -> bytes:
        doc = Document()
        for line in content_markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            m = _HEADING_RE.match(stripped)
            if m:
                level = min(len(m.group(1)), 9)
                doc.add_heading(m.group(2), level=level)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def filename(self, title: str) -> str:
        return f"{sanitize_filename(title)}.docx"
