"""End-to-end extraction tests for the 7b-i formats, using real files
constructed in-test: .md/.txt/.csv (bytes), .pdf (handcrafted minimal PDF),
.docx (python-docx), .xlsx (openpyxl)."""

from __future__ import annotations

import io

from skylize.memory.extraction import extract_text

def _minimal_pdf(text: str) -> bytes:
    """Build a fully valid single-page PDF (with xref + trailer) by hand."""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    ).encode()
    return bytes(out)


def test_extract_markdown() -> None:
    text = extract_text("notes.md", b"# Ops Notes\n\nDaily standup at 9.")
    assert "Ops Notes" in text and "Daily standup" in text


def test_extract_txt() -> None:
    text = extract_text("memo.txt", "Simple memo with unicode \xe9\xe0.".encode("utf-8"))
    assert "Simple memo" in text


def test_extract_csv() -> None:
    text = extract_text("tasks.csv", b"task,due\nInvoice review,2026-07-01\n")
    assert "Invoice review" in text


def test_extract_pdf() -> None:
    text = extract_text("policy.pdf", _minimal_pdf("Overdue tasks summary policy"))
    assert "Overdue tasks summary policy" in text


def test_extract_docx() -> None:
    import docx

    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Quarterly marketing campaign brief.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Channel"
    table.rows[0].cells[1].text = "Paid social"
    document.save(buf)

    text = extract_text("brief.docx", buf.getvalue())
    assert "Quarterly marketing campaign brief." in text
    assert "Channel | Paid social" in text


def test_extract_xlsx() -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws.append(["Line", "Amount"])
    ws.append(["Ad spend", 12000])
    buf = io.BytesIO()
    wb.save(buf)

    text = extract_text("budget.xlsx", buf.getvalue())
    assert "# Budget" in text
    assert "Ad spend | 12000" in text
